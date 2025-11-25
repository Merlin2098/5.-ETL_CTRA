"""
Módulo de validación para generación de certificados.
Valida archivos, datos y configuraciones antes de la generación.
"""

import os
import re
from pathlib import Path
from typing import Tuple, List
import pandas as pd
from docx import Document


class CertificateValidator:
    """Valida archivos, datos y configuraciones antes de generar certificados"""
    
    # Columnas requeridas en el DataFrame
    REQUIRED_COLUMNS = [
        'APELLIDOS Y NOMBRES',
        'FECHAS_CERTIFICADO',
        'DÍAS_LABORADOS',
        'CARGO',
        'CLIENTE',
        'FECHA_GENERAR'
    ]
    
    # Placeholders requeridos en la plantilla
    REQUIRED_PLACEHOLDERS = [
        'APELLIDOS_Y_NOMBRES',
        'FECHAS_CERTIFICADO',
        'DÍAS_LABORADOS',
        'CARGO',
        'CLIENTE',
        'FECHA_GENERAR'
    ]
    
    @staticmethod
    def validate_template_exists(template_path: str) -> Tuple[bool, str]:
        """
        Valida que la plantilla Word existe y es accesible.
        
        Args:
            template_path: Ruta a la plantilla Word
            
        Returns:
            Tupla (éxito, mensaje_error)
        """
        # Validar que la ruta no esté vacía
        if not template_path or not template_path.strip():
            return False, "La ruta de la plantilla está vacía"
        
        # Validar extensión
        if not template_path.lower().endswith('.docx'):
            return False, "La plantilla debe ser un archivo .docx"
        
        # Validar existencia
        if not os.path.exists(template_path):
            return False, f"La plantilla no existe: {template_path}"
        
        # Validar que es un archivo (no directorio)
        if not os.path.isfile(template_path):
            return False, f"La ruta no es un archivo válido: {template_path}"
        
        # Intentar abrir con python-docx
        try:
            doc = Document(template_path)
            # Verificar que tiene al menos un párrafo
            if not doc.paragraphs:
                return False, "La plantilla está vacía o corrupta"
        except Exception as e:
            return False, f"Error al abrir la plantilla: {str(e)}"
        
        return True, ""
    
    @staticmethod
    def validate_template_placeholders(template_path: str) -> Tuple[bool, List[str]]:
        """
        Extrae y valida placeholders de la plantilla.
        
        Args:
            template_path: Ruta a la plantilla Word
            
        Returns:
            Tupla (éxito, lista_placeholders_o_mensaje_error)
        """
        try:
            doc = Document(template_path)
            
            # Buscar todos los placeholders {{...}}
            placeholders = set()
            pattern = r'\{\{([^}]+)\}\}'
            
            for para in doc.paragraphs:
                matches = re.findall(pattern, para.text)
                for match in matches:
                    # Limpiar espacios
                    placeholders.add(match.strip())
            
            # Convertir a lista ordenada
            placeholders_list = sorted(list(placeholders))
            
            # Validar que estén todos los requeridos
            missing = []
            for required in CertificateValidator.REQUIRED_PLACEHOLDERS:
                if required not in placeholders:
                    missing.append(required)
            
            if missing:
                error_msg = f"Faltan placeholders requeridos: {', '.join(missing)}"
                return False, [error_msg]
            
            return True, placeholders_list
            
        except Exception as e:
            return False, [f"Error al extraer placeholders: {str(e)}"]
    
    @staticmethod
    def validate_dataframe(df: pd.DataFrame) -> Tuple[bool, str]:
        """
        Valida integridad del DataFrame.
        
        Args:
            df: DataFrame a validar
            
        Returns:
            Tupla (éxito, mensaje_error)
        """
        # Verificar que no sea None
        if df is None:
            return False, "El DataFrame es None"
        
        # Verificar que no esté vacío
        if df.empty:
            return False, "El DataFrame está vacío"
        
        # Verificar columnas requeridas
        missing_cols = []
        for col in CertificateValidator.REQUIRED_COLUMNS:
            if col not in df.columns:
                missing_cols.append(col)
        
        if missing_cols:
            return False, f"Faltan columnas requeridas: {', '.join(missing_cols)}"
        
        # Verificar que no haya nulos en columnas críticas
        null_counts = {}
        for col in CertificateValidator.REQUIRED_COLUMNS:
            null_count = df[col].isnull().sum()
            if null_count > 0:
                null_counts[col] = null_count
        
        if null_counts:
            error_parts = [f"{col}: {count} nulos" for col, count in null_counts.items()]
            return False, f"Se encontraron valores nulos: {', '.join(error_parts)}"
        
        # Verificar que DÍAS_LABORADOS sea numérico
        if 'DÍAS_LABORADOS' in df.columns:
            if not pd.api.types.is_numeric_dtype(df['DÍAS_LABORADOS']):
                return False, "La columna DÍAS_LABORADOS debe ser numérica"
        
        return True, ""
    
    @staticmethod
    def validate_row_data(row: pd.Series) -> Tuple[bool, str]:
        """
        Valida que una fila tenga todos los datos necesarios.
        
        Args:
            row: Fila del DataFrame a validar
            
        Returns:
            Tupla (éxito, mensaje_error)
        """
        # Verificar cada campo requerido
        for col in CertificateValidator.REQUIRED_COLUMNS:
            if col not in row.index:
                return False, f"Falta columna: {col}"
            
            value = row[col]
            
            # Verificar que no sea nulo
            if pd.isnull(value):
                return False, f"Campo vacío: {col}"
            
            # Verificar que no sea string vacío (para campos de texto)
            if isinstance(value, str) and not value.strip():
                return False, f"Campo vacío: {col}"
        
        # Verificar específicamente APELLIDOS Y NOMBRES
        apellidos = row.get('APELLIDOS Y NOMBRES', '')
        if not apellidos or not str(apellidos).strip():
            return False, "APELLIDOS Y NOMBRES está vacío"
        
        # Verificar que DÍAS_LABORADOS sea positivo
        dias = row.get('DÍAS_LABORADOS', 0)
        try:
            dias_num = int(dias)
            if dias_num <= 0:
                return False, "DÍAS_LABORADOS debe ser mayor a 0"
        except (ValueError, TypeError):
            return False, "DÍAS_LABORADOS debe ser un número válido"
        
        return True, ""
    
    @staticmethod
    def validate_output_folder(folder_path: str) -> Tuple[bool, str]:
        """
        Valida que la carpeta de salida sea válida.
        
        Args:
            folder_path: Ruta de la carpeta de salida
            
        Returns:
            Tupla (éxito, mensaje_error)
        """
        # Validar que la ruta no esté vacía
        if not folder_path or not folder_path.strip():
            return False, "La ruta de la carpeta está vacía"
        
        try:
            # Convertir a Path para mejor manejo
            path = Path(folder_path)
            
            # Si no existe, intentar crear
            if not path.exists():
                path.mkdir(parents=True, exist_ok=True)
            
            # Verificar que es un directorio
            if not path.is_dir():
                return False, f"La ruta no es un directorio válido: {folder_path}"
            
            # Verificar permisos de escritura
            test_file = path / ".write_test"
            try:
                test_file.touch()
                test_file.unlink()
            except Exception as e:
                return False, f"No hay permisos de escritura en: {folder_path}"
            
            return True, ""
            
        except Exception as e:
            return False, f"Error al validar carpeta: {str(e)}"
    
    @staticmethod
    def validate_filename(apellidos_nombres: str, max_length: int = 100) -> str:
        """
        Sanitiza el nombre de archivo eliminando caracteres no válidos.
        
        Args:
            apellidos_nombres: Nombre a sanitizar
            max_length: Longitud máxima del nombre
            
        Returns:
            Nombre sanitizado
        """
        if not apellidos_nombres:
            return "SIN_NOMBRE"
        
        # Convertir a string por si acaso
        nombre = str(apellidos_nombres).strip()
        
        # Eliminar caracteres no válidos para nombres de archivo
        # Caracteres prohibidos: / \ : * ? " < > |
        invalid_chars = r'[/\\:*?"<>|]'
        nombre = re.sub(invalid_chars, '', nombre)
        
        # Reemplazar espacios múltiples por uno solo
        nombre = re.sub(r'\s+', ' ', nombre)
        
        # Eliminar espacios al inicio y final
        nombre = nombre.strip()
        
        # Si quedó vacío después de sanitizar
        if not nombre:
            return "SIN_NOMBRE"
        
        # Truncar si es muy largo
        if len(nombre) > max_length:
            nombre = nombre[:max_length].strip()
        
        return nombre