"""
Módulo de generación de certificados Word desde plantilla.
Genera certificados individuales y en batch con reemplazo de placeholders.
"""

import os
import re
from pathlib import Path
from typing import Tuple, Dict, List, Optional, Callable, Any
import pandas as pd
from docx import Document


class WordCertificateGenerator:
    """Generador de certificados Word desde plantilla"""
    
    def __init__(self, template_path: str):
        """
        Inicializa el generador con una plantilla.
        
        Args:
            template_path: Ruta a la plantilla Word
            
        Raises:
            ValueError: Si la plantilla no existe o no es válida
        """
        self.template_path = template_path
        self.template = None
        self._load_template()
    
    def _load_template(self) -> None:
        """Carga la plantilla Word en memoria"""
        if not os.path.exists(self.template_path):
            raise ValueError(f"La plantilla no existe: {self.template_path}")
        
        try:
            self.template = Document(self.template_path)
        except Exception as e:
            raise ValueError(f"Error al cargar plantilla: {str(e)}")
    
    def get_placeholders(self) -> List[str]:
        """
        Extrae placeholders de la plantilla actual.
        
        Returns:
            Lista de placeholders únicos encontrados (sin llaves)
        """
        placeholders = set()
        pattern = r'\{\{([^}]+)\}\}'
        
        for para in self.template.paragraphs:
            matches = re.findall(pattern, para.text)
            for match in matches:
                placeholders.add(match.strip())
        
        return sorted(list(placeholders))
    
    def _replace_placeholders(self, doc: Document, data: dict) -> Document:
        """
        Reemplaza placeholders en un documento.
        
        Args:
            doc: Documento Word
            data: Diccionario con datos para reemplazar
            
        Returns:
            Documento modificado
        """
        # Iterar sobre todos los párrafos
        for para in doc.paragraphs:
            # Obtener el texto original
            original_text = para.text
            modified_text = original_text
            
            # Buscar y reemplazar placeholders
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in modified_text:
                    # Convertir valor a string
                    str_value = str(value) if value is not None else ""
                    modified_text = modified_text.replace(placeholder, str_value)
            
            # Si el texto cambió, reemplazar el párrafo completo
            if modified_text != original_text:
                # Limpiar runs existentes
                for run in para.runs:
                    run.text = ""
                # Agregar nuevo texto al primer run
                if para.runs:
                    para.runs[0].text = modified_text
                else:
                    para.add_run(modified_text)
        
        return doc
    
    def _prepare_data_dict(self, row_data: pd.Series) -> dict:
        """
        Prepara diccionario de datos desde una fila del DataFrame.
        Convierte nombres de columnas al formato de placeholders.
        
        Args:
            row_data: Fila del DataFrame
            
        Returns:
            Diccionario con datos preparados
        """
        data = {}
        
        # Mapeo de columnas a placeholders
        column_mapping = {
            'APELLIDOS Y NOMBRES': 'APELLIDOS_Y_NOMBRES',
            'FECHAS_CERTIFICADO': 'FECHAS_CERTIFICADO',
            'DÍAS_LABORADOS': 'DÍAS_LABORADOS',
            'CARGO': 'CARGO',
            'CLIENTE': 'CLIENTE',
            'FECHA_GENERAR': 'FECHA_GENERAR'
        }
        
        # Convertir datos
        for col_name, placeholder_name in column_mapping.items():
            if col_name in row_data.index:
                value = row_data[col_name]
                data[placeholder_name] = value
            else:
                data[placeholder_name] = ""
        
        return data
    
    def _sanitize_filename(self, apellidos_nombres: str, max_length: int = 100) -> str:
        """
        Sanitiza nombre de archivo.
        
        Args:
            apellidos_nombres: Nombre para el archivo
            max_length: Longitud máxima
            
        Returns:
            Nombre sanitizado
        """
        if not apellidos_nombres:
            return "SIN_NOMBRE"
        
        nombre = str(apellidos_nombres).strip()
        
        # Eliminar caracteres no válidos
        invalid_chars = r'[/\\:*?"<>|]'
        nombre = re.sub(invalid_chars, '', nombre)
        
        # Reemplazar espacios múltiples
        nombre = re.sub(r'\s+', ' ', nombre)
        
        nombre = nombre.strip()
        
        if not nombre:
            return "SIN_NOMBRE"
        
        # Truncar si es necesario
        if len(nombre) > max_length:
            nombre = nombre[:max_length].strip()
        
        return nombre
    
    def generate_single(self, 
                       row_data: pd.Series, 
                       output_path: str) -> Tuple[bool, str]:
        """
        Genera un certificado individual.
        
        Args:
            row_data: Fila con datos del certificado
            output_path: Ruta de salida completa del archivo
            
        Returns:
            Tupla (éxito, mensaje)
        """
        try:
            # Crear una copia de la plantilla
            doc = Document(self.template_path)
            
            # Preparar datos
            data_dict = self._prepare_data_dict(row_data)
            
            # Reemplazar placeholders
            doc = self._replace_placeholders(doc, data_dict)
            
            # Asegurar que el directorio existe
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # Guardar documento
            doc.save(output_path)
            
            return True, output_path
            
        except Exception as e:
            return False, f"Error al generar certificado: {str(e)}"
    
    def _handle_duplicate_filename(self, output_folder: str, base_filename: str) -> str:
        """
        Maneja nombres de archivo duplicados agregando sufijos.
        
        Args:
            output_folder: Carpeta de salida
            base_filename: Nombre base del archivo (sin extensión)
            
        Returns:
            Nombre de archivo único
        """
        output_path = os.path.join(output_folder, f"{base_filename}.docx")
        
        # Si no existe, usar tal cual
        if not os.path.exists(output_path):
            return output_path
        
        # Si existe, agregar sufijo numérico
        counter = 1
        while True:
            new_filename = f"{base_filename}_{counter}.docx"
            output_path = os.path.join(output_folder, new_filename)
            
            if not os.path.exists(output_path):
                return output_path
            
            counter += 1
            
            # Límite de seguridad
            if counter > 1000:
                raise ValueError(f"Demasiados duplicados para: {base_filename}")
    
    def generate_batch(self,
                      df: pd.DataFrame,
                      output_folder: str,
                      progress_callback: Optional[Callable[[int, str, int, int], None]] = None
                      ) -> Dict[str, Any]:
        """
        Genera múltiples certificados en batch.
        
        Args:
            df: DataFrame con datos de certificados
            output_folder: Carpeta de salida
            progress_callback: Función callback(índice, mensaje, actual, total)
            
        Returns:
            Diccionario con resultados:
            {
                'total': int,
                'exitosos': int,
                'fallidos': int,
                'archivos_generados': List[str],
                'errores': List[Dict]
            }
        """
        # Asegurar que la carpeta existe
        os.makedirs(output_folder, exist_ok=True)
        
        # Inicializar resultados
        results = {
            'total': len(df),
            'exitosos': 0,
            'fallidos': 0,
            'archivos_generados': [],
            'errores': []
        }
        
        # Procesar cada fila
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            try:
                # Obtener apellidos y nombres para el archivo
                apellidos_nombres = row.get('APELLIDOS Y NOMBRES', 'SIN_NOMBRE')
                
                # Sanitizar nombre
                nombre_sanitizado = self._sanitize_filename(apellidos_nombres)
                
                # Generar nombre de archivo con manejo de duplicados
                base_filename = f"Certificado_{nombre_sanitizado}"
                output_path = self._handle_duplicate_filename(output_folder, base_filename)
                
                # Generar certificado
                success, message = self.generate_single(row, output_path)
                
                if success:
                    results['exitosos'] += 1
                    results['archivos_generados'].append(output_path)
                    
                    # Callback de progreso
                    if progress_callback:
                        progress_callback(
                            idx,
                            f"Generado: {os.path.basename(output_path)}",
                            results['exitosos'],
                            results['total']
                        )
                else:
                    results['fallidos'] += 1
                    results['errores'].append({
                        'fila': idx,
                        'apellidos': apellidos_nombres,
                        'error': message
                    })
                    
                    # Callback de progreso con error
                    if progress_callback:
                        progress_callback(
                            idx,
                            f"ERROR: {apellidos_nombres} - {message}",
                            results['exitosos'],
                            results['total']
                        )
                
            except Exception as e:
                results['fallidos'] += 1
                apellidos = row.get('APELLIDOS Y NOMBRES', 'DESCONOCIDO')
                error_msg = f"Excepción: {str(e)}"
                
                results['errores'].append({
                    'fila': idx,
                    'apellidos': apellidos,
                    'error': error_msg
                })
                
                # Callback de progreso con excepción
                if progress_callback:
                    progress_callback(
                        idx,
                        f"EXCEPCIÓN: {apellidos} - {error_msg}",
                        results['exitosos'],
                        results['total']
                    )
        
        return results